"""
ULTRON AI – Text Extraction Utility
Extracts plain text from PDF, DOCX, and TXT files.
Uses PyMuPDF (fitz) as primary PDF parser with pdfplumber fallback.
"""

import os
import io
import chardet
from pathlib import Path


def extract_text_from_file(file_path: str) -> str:
    """
    Extract plain text from a resume file.
    Supports: .pdf, .docx, .txt
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext == ".txt":
        return _extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Extract text from in-memory file bytes."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf_bytes(file_bytes)
    elif ext == ".docx":
        return _extract_docx_bytes(file_bytes)
    elif ext == ".txt":
        encoding = chardet.detect(file_bytes).get("encoding", "utf-8") or "utf-8"
        return file_bytes.decode(encoding, errors="replace")
    else:
        raise ValueError(f"Unsupported file format: {ext}")


# ------------------------------------------------------------------
# PDF Extraction
# ------------------------------------------------------------------

def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF with pdfplumber fallback."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        if text.strip():
            return text
    except Exception:
        pass

    # Fallback to pdfplumber
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        return text
    except Exception as e:
        raise RuntimeError(f"Failed to extract PDF text: {e}")


def _extract_pdf_bytes(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF with pdfplumber fallback."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        if text.strip():
            return text
    except Exception:
        pass

    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        return text
    except Exception as e:
        raise RuntimeError(f"Failed to extract PDF bytes: {e}")


# ------------------------------------------------------------------
# DOCX Extraction
# ------------------------------------------------------------------

def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"Failed to extract DOCX text: {e}")


def _extract_docx_bytes(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"Failed to extract DOCX bytes: {e}")


# ------------------------------------------------------------------
# TXT Extraction
# ------------------------------------------------------------------

def _extract_txt(file_path: str) -> str:
    """Extract text from TXT file with encoding detection."""
    with open(file_path, "rb") as f:
        raw = f.read()
    encoding = chardet.detect(raw).get("encoding", "utf-8") or "utf-8"
    return raw.decode(encoding, errors="replace")
